import os, io, sys
import fitz  # PyMuPDF
from PIL import Image
from pillow_heif import register_heif_opener
from docx import Document
from docx.shared import Inches
from django.conf import settings
from datetime import datetime

# Register HEIF opener with Pillow
register_heif_opener()

def get_unique_filename(filename, ext):
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    name = os.path.splitext(filename)[0]
    return f"{name}_{timestamp}.{ext}"

def convert_image(input_file, target_format):
    """
    Converts an image to the target format (JPG or PNG).
    Handles HEIC automatically.
    """
    image = Image.open(input_file)
    
    # If target is JPEG and image has alpha channel, convert to RGB
    if target_format.upper() in ['JPG', 'JPEG'] and image.mode in ('RGBA', 'P', 'LA'):
        image = image.convert('RGB')
    elif image.mode == 'RGBA' and target_format.upper() == 'PNG':
        pass # Keep RGBA for PNG
    else:
        image = image.convert('RGB')

    filename = get_unique_filename(input_file.name, target_format.lower())
    relative_path = f"conversions/{filename}"
    absolute_path = os.path.join(settings.MEDIA_ROOT, relative_path)
    
    os.makedirs(os.path.dirname(absolute_path), exist_ok=True)
    
    image.save(absolute_path, target_format.upper() if target_format.upper() != 'JPG' else 'JPEG')
    return relative_path

def convert_to_pdf(input_file):
    """
    Converts an image (including HEIC) to PDF.
    """
    image = Image.open(input_file)
    
    # PDF requires RGB mode
    if image.mode != 'RGB':
        image = image.convert('RGB')
        
    filename = get_unique_filename(input_file.name, 'pdf')
    relative_path = f"pdfs/{filename}"
    absolute_path = os.path.join(settings.MEDIA_ROOT, relative_path)
    
    os.makedirs(os.path.dirname(absolute_path), exist_ok=True)
    
    image.save(absolute_path, "PDF")
    return relative_path

def compress_image(input_file, quality=60):
    """
    Compresses an image by trying both JPEG and PNG optimisation,
    then returns whichever output is smallest.
    If neither beats the original, the original file is returned as-is.
    quality: integer 1-95 (lower = smaller JPEG)
    Returns: (relative_path, original_size_bytes, compressed_size_bytes, already_optimal)
    """
    input_file.seek(0)
    original_data = input_file.read()
    original_size = len(original_data)

    img = Image.open(io.BytesIO(original_data))
    img.load()

    base_name = os.path.splitext(input_file.name)[0]
    output_dir = os.path.join(settings.MEDIA_ROOT, 'compressed')
    os.makedirs(output_dir, exist_ok=True)

    candidates = []  # list of (size_bytes, ext, buffer)

    # --- Candidate 1: JPEG at requested quality ---
    rgb_img = img.copy()
    if rgb_img.mode in ('RGBA', 'P', 'LA'):
        background = Image.new('RGB', rgb_img.size, (255, 255, 255))
        if rgb_img.mode == 'RGBA':
            background.paste(rgb_img, mask=rgb_img.split()[3])
        else:
            background.paste(rgb_img.convert('RGBA'), mask=rgb_img.convert('RGBA').split()[3])
        rgb_img = background
    elif rgb_img.mode != 'RGB':
        rgb_img = rgb_img.convert('RGB')

    jpeg_buf = io.BytesIO()
    rgb_img.save(jpeg_buf, 'JPEG', quality=quality, optimize=True)
    candidates.append((jpeg_buf.tell(), 'jpg', jpeg_buf.getvalue()))

    # --- Candidate 2: PNG with optimize flag ---
    png_buf = io.BytesIO()
    png_img = img.copy()
    png_img.save(png_buf, 'PNG', optimize=True, compress_level=9)
    candidates.append((png_buf.tell(), 'png', png_buf.getvalue()))

    # Pick the smallest candidate
    best_size, best_ext, best_data = min(candidates, key=lambda x: x[0])

    if best_size >= original_size:
        # Nothing helped — return original file unchanged
        orig_ext = os.path.splitext(input_file.name)[1].lstrip('.') or 'jpg'
        filename = get_unique_filename(base_name, orig_ext)
        relative_path = f"compressed/{filename}"
        absolute_path = os.path.join(settings.MEDIA_ROOT, relative_path)
        with open(absolute_path, 'wb') as f:
            f.write(original_data)
        return relative_path, original_size, original_size

    filename = get_unique_filename(base_name, best_ext)
    relative_path = f"compressed/{filename}"
    absolute_path = os.path.join(settings.MEDIA_ROOT, relative_path)
    with open(absolute_path, 'wb') as f:
        f.write(best_data)

    return relative_path, original_size, best_size


def compress_pdf(input_file, image_quality=60):
    """
    Compresses a PDF by re-encoding all embedded images at a lower JPEG quality.
    image_quality: 1-95 (lower = smaller output)
    Returns: (relative_path, original_size_bytes, compressed_size_bytes)
    """
    input_file.seek(0)
    original_data = input_file.read()
    original_size = len(original_data)

    doc = fitz.open(stream=original_data, filetype='pdf')

    for page in doc:
        # Iterate over all images on this page
        for img_info in page.get_images(full=True):
            xref = img_info[0]
            try:
                base_image = doc.extract_image(xref)
                img_bytes = base_image['image']
                img_ext = base_image['ext']

                # Re-compress supported formats
                if img_ext.lower() in ('jpeg', 'jpg', 'png', 'bmp', 'tiff'):
                    # Check current filter to avoid re-compressing already optimal JEPGs
                    img_obj_dict = doc.xref_get_keys(xref)
                    is_already_dct = "Filter" in img_obj_dict and "/DCTDecode" in doc.xref_get_value(xref)

                    pil_img = Image.open(io.BytesIO(img_bytes))
                    
                    # Convert to RGB (flatten transparency as JPEG doesn't support it)
                    if pil_img.mode in ('RGBA', 'P', 'LA'):
                        bg = Image.new('RGB', pil_img.size, (255, 255, 255))
                        if pil_img.mode == 'RGBA':
                            bg.paste(pil_img, mask=pil_img.split()[3])
                        else:
                            bg.paste(pil_img.convert('RGB'))
                        pil_img = bg
                    elif pil_img.mode != 'RGB':
                        pil_img = pil_img.convert('RGB')

                    buf = io.BytesIO()
                    pil_img.save(buf, 'JPEG', quality=image_quality, optimize=True)
                    new_img_bytes = buf.getvalue()

                    # Only replace if the recompressed version is significantly smaller
                    if len(new_img_bytes) < len(img_bytes) * 0.95:
                        doc.update_stream(xref, new_img_bytes)
                        # IMPORTANT: Must update the metadata so the PDF reader knows it's now a JPEG
                        doc.xref_set_key(xref, "Filter", "/DCTDecode")
                        doc.xref_set_key(xref, "ColorSpace", "/DeviceRGB")
            except Exception:
                continue  # Skip unprocessable images

    # Save compressed PDF to buffer
    compressed_buf = io.BytesIO()
    doc.save(
        compressed_buf,
        garbage=4,       # remove unused objects
        deflate=True,    # compress streams
        clean=True,
    )
    doc.close()
    compressed_data = compressed_buf.getvalue()
    compressed_size = len(compressed_data)

    base_name = os.path.splitext(input_file.name)[0]
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    filename = f"{base_name}_{timestamp}_compressed.pdf"
    relative_path = f"compressed/{filename}"
    absolute_path = os.path.join(settings.MEDIA_ROOT, relative_path)
    os.makedirs(os.path.dirname(absolute_path), exist_ok=True)

    if compressed_size < original_size:
        with open(absolute_path, 'wb') as f:
            f.write(compressed_data)
    else:
        # Compression didn't help — save original
        with open(absolute_path, 'wb') as f:
            f.write(original_data)
        compressed_size = original_size

    return relative_path, original_size, compressed_size

def merge_images_to_pdf(input_files):
    """
    Merges multiple images into a single multi-page PDF.
    Uses BytesIO to ensure file data is independent of original pointers.
    """
    sys.stderr.write(f"\n[MERGE DEBUG] Received {len(input_files)} files\n")
    processed_images = []
    
    for i, f in enumerate(input_files):
        try:
            f.seek(0)
            file_data = f.read()
            sys.stderr.write(f"[MERGE DEBUG] File {i+1} size: {len(file_data)} bytes\n")
            
            img = Image.open(io.BytesIO(file_data))
            img.load()
            
            # Create a clean RGB copy
            rgb_img = Image.new("RGB", img.size, (255, 255, 255))
            if img.mode == 'RGBA':
                rgb_img.paste(img, mask=img.split()[3])
            else:
                rgb_img.paste(img)
            
            processed_images.append(rgb_img)
            sys.stderr.write(f"[MERGE DEBUG] Image {i+1} processed: {img.size} {img.mode}\n")
        except Exception as e:
            sys.stderr.write(f"[MERGE DEBUG] ERROR on file {i+1}: {str(e)}\n")
    
    if len(processed_images) < 1:
        sys.stderr.write("[MERGE DEBUG] No images to merge\n")
        return None
        
    filename = get_unique_filename("merged", 'pdf')
    relative_path = f"pdfs/{filename}"
    absolute_path = os.path.join(settings.MEDIA_ROOT, relative_path)
    
    os.makedirs(os.path.dirname(absolute_path), exist_ok=True)
    
    try:
        if len(processed_images) > 1:
            # save_all=True and append_images are the keys for multi-page
            processed_images[0].save(
                absolute_path, 
                "PDF", 
                save_all=True, 
                append_images=processed_images[1:],
                resolution=100.0,
                quality=95
            )
            sys.stderr.write(f"[MERGE DEBUG] SUCCESS: Saved {len(processed_images)} pages to {relative_path}\n")
        else:
            processed_images[0].save(absolute_path, "PDF")
            sys.stderr.write(f"[MERGE DEBUG] SUCCESS: Saved 1 page to {relative_path}\n")
    except Exception as e:
        sys.stderr.write(f"[MERGE DEBUG] SAVE ERROR: {str(e)}\n")
        
    for img in processed_images:
        img.close()
        
    return relative_path

def convert_to_word(input_file):
    """
    Converts an image (including HEIC) to a Word document.
    """
    # First convert to a temporary image that Pillow/python-docx can definitely handle
    image = Image.open(input_file)
    if image.mode != 'RGB':
        image = image.convert('RGB')
    
    temp_img_path = os.path.join(settings.MEDIA_ROOT, 'temp', f"temp_{input_file.name}.jpg")
    os.makedirs(os.path.dirname(temp_img_path), exist_ok=True)
    image.save(temp_img_path, "JPEG")
    
    # Create Word document
    doc = Document()
    doc.add_heading('Converted Image', 0)
    
    # Add image to document
    doc.add_picture(temp_img_path, width=Inches(6))
    
    filename = get_unique_filename(input_file.name, 'docx')
    relative_path = f"word_docs/{filename}"
    absolute_path = os.path.join(settings.MEDIA_ROOT, relative_path)
    
    os.makedirs(os.path.dirname(absolute_path), exist_ok=True)
    doc.save(absolute_path)
    
    # Clean up temp image
    if os.path.exists(temp_img_path):
        os.remove(temp_img_path)
        
    return relative_path
